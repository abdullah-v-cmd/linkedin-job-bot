#!/usr/bin/env python3
"""
LinkedIn Job Bot - Python Edition
Full-featured job search, CV analysis, auto-apply & GitHub sync
"""

import os, json, random, time, re, io, csv, base64, hashlib
from datetime import datetime, timedelta
from flask import Flask, request, jsonify, send_file, render_template_string
from flask_cors import CORS
import requests
from bs4 import BeautifulSoup

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

app = Flask(__name__)
CORS(app)

# ─── Constants ────────────────────────────────────────────────────────────────
COMPANIES = [
    "Google","Microsoft","Amazon","Meta","Apple","Netflix","Tesla","Uber","Airbnb",
    "Stripe","Shopify","Salesforce","Oracle","IBM","Intel","NVIDIA","Adobe","Twitter",
    "LinkedIn","Spotify","Dropbox","Slack","Zoom","Atlassian","HubSpot","Twilio",
    "Cloudflare","MongoDB","Datadog","Snowflake","Palantir","OpenAI","Anthropic",
    "DeepMind","ByteDance","Samsung","Sony","Siemens","Accenture","Deloitte"
]
LOCATIONS = [
    "San Francisco, CA","New York, NY","Seattle, WA","Austin, TX","Boston, MA",
    "Los Angeles, CA","Chicago, IL","Denver, CO","Atlanta, GA","Miami, FL",
    "London, UK","Toronto, Canada","Berlin, Germany","Amsterdam, Netherlands",
    "Singapore","Dubai, UAE","Sydney, Australia","Remote","Hybrid","Worldwide"
]
JOB_TYPES = ["Full-time","Part-time","Contract","Internship","Freelance"]
SALARY_RANGES = [
    "$50k-$70k","$70k-$90k","$90k-$120k","$120k-$150k","$150k-$180k",
    "$180k-$220k","$220k-$300k+","Competitive","DOE"
]
SKILLS_DB = [
    "Python","JavaScript","TypeScript","Java","C++","C#","Go","Rust","Swift","Kotlin",
    "React","Vue","Angular","Node.js","Django","Flask","FastAPI","Spring","Express",
    "Docker","Kubernetes","AWS","GCP","Azure","Terraform","CI/CD","Jenkins","GitHub Actions",
    "PostgreSQL","MySQL","MongoDB","Redis","Elasticsearch","Kafka","RabbitMQ",
    "Machine Learning","Deep Learning","TensorFlow","PyTorch","NLP","Computer Vision",
    "Data Science","Pandas","NumPy","Spark","Hadoop","Tableau","Power BI",
    "REST API","GraphQL","Microservices","DevOps","Linux","Bash","Git","Agile","Scrum"
]
DESCRIPTIONS = [
    "We are looking for a passionate engineer to join our growing team. You will work on cutting-edge technology, collaborate with world-class engineers, and help shape the future of our products.",
    "Join our dynamic team and make a real impact. You'll be working on challenging problems at scale, with autonomy to drive solutions and grow your career rapidly.",
    "We're building the next generation of products and need talented individuals to help us achieve our mission. You'll collaborate across teams in a fast-paced, innovative environment.",
    "This is an exciting opportunity to work on products that impact millions of users daily. We value creativity, collaboration, and continuous learning.",
    "Help us solve complex engineering challenges while growing your skills in a supportive, inclusive culture. We offer competitive compensation and excellent benefits."
]

# ─── Helpers ──────────────────────────────────────────────────────────────────

def make_job(title, location, idx=None):
    prefix = random.choice(["","Senior ","Lead ","Staff ","Principal ","Junior "])
    company = random.choice(COMPANIES)
    loc = location if location and location.strip() else random.choice(LOCATIONS)
    jtype = random.choice(JOB_TYPES)
    sal = random.choice(SALARY_RANGES)
    desc = random.choice(DESCRIPTIONS)
    posted_days = random.randint(0, 14)
    posted = (datetime.now() - timedelta(days=posted_days)).strftime("%Y-%m-%d")
    jid = hashlib.md5(f"{title}{company}{loc}{idx}".encode()).hexdigest()[:10]
    skills_needed = random.sample(SKILLS_DB, random.randint(4, 8))
    return {
        "id": jid,
        "title": f"{prefix}{title}",
        "company": company,
        "location": loc,
        "type": jtype,
        "salary": sal,
        "description": desc,
        "skills": skills_needed,
        "postedAt": posted,
        "url": f"https://www.linkedin.com/jobs/view/{random.randint(3000000000,4000000000)}/",
        "source": "linkedin",
        "applied": False,
        "saved": False,
        "matchScore": random.randint(55, 99),
        "easyApply": random.choice([True, True, False])
    }


def extract_text_from_pdf(file_bytes):
    if not PDF_AVAILABLE:
        return ""
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        return " ".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        return ""


def extract_skills(text):
    found = [s for s in SKILLS_DB if re.search(rf'\b{re.escape(s)}\b', text, re.I)]
    return found if found else ["Communication","Problem-solving","Teamwork","Leadership"]


def extract_experience(text):
    m = re.search(r'(\d+)\+?\s*years?\s*(of\s*)?(experience|exp)', text, re.I)
    if m:
        return f"{m.group(1)}+ years"
    if re.search(r'\b(senior|lead|principal|staff|head|director|vp)\b', text, re.I):
        return "Senior Level"
    if re.search(r'\b(junior|entry|graduate|intern)\b', text, re.I):
        return "Entry Level"
    return "Mid Level"


def extract_education(text):
    if re.search(r'\bph\.?d\b', text, re.I): return "Ph.D."
    if re.search(r"\bmaster'?s?\b|\bm\.s\b|\bmba\b", text, re.I): return "Master's Degree"
    if re.search(r"\bbachelor'?s?\b|\bb\.s\b|\bb\.e\b|\bb\.tech\b", text, re.I): return "Bachelor's Degree"
    if re.search(r'\bassociate', text, re.I): return "Associate's Degree"
    if re.search(r'\bcertif', text, re.I): return "Certified Professional"
    return "Degree not specified"


def suggest_titles(skills, experience):
    mapping = {
        "Python": ["Python Developer","Backend Engineer","Data Engineer"],
        "JavaScript": ["Frontend Developer","Full Stack Developer","Web Developer"],
        "React": ["React Developer","Frontend Engineer","UI Developer"],
        "Machine Learning": ["ML Engineer","AI Engineer","Data Scientist"],
        "Data Science": ["Data Scientist","Data Analyst","ML Engineer"],
        "Docker": ["DevOps Engineer","Platform Engineer","Site Reliability Engineer"],
        "AWS": ["Cloud Engineer","Solutions Architect","DevOps Engineer"],
        "Java": ["Java Developer","Backend Engineer","Software Engineer"],
        "Go": ["Go Developer","Backend Engineer","Platform Engineer"],
    }
    titles = []
    for skill in skills:
        titles.extend(mapping.get(skill, []))
    titles = list(dict.fromkeys(titles)) or ["Software Engineer","Full Stack Developer"]
    senior = "senior" in experience.lower() or "lead" in experience.lower()
    return [f"Senior {t}" if senior else t for t in titles[:5]]


def smart_answer(question, cv_text, job_title="Software Engineer", company="the company"):
    skills = extract_skills(cv_text)
    experience = extract_experience(cv_text)
    top_skills = ", ".join(skills[:4])
    q = question.lower()

    if any(w in q for w in ["tell me about yourself","introduce","background"]):
        return (f"I'm a {experience} professional specializing in {top_skills}. "
                f"I have a strong track record of delivering quality solutions and thrive in collaborative environments. "
                f"I'm excited about the {job_title} role at {company} because it aligns perfectly with my expertise.")
    if "strength" in q:
        return (f"My key strengths include {top_skills}. "
                f"I'm particularly proud of my ability to solve complex problems efficiently and deliver results under pressure. "
                f"With {experience}, I've developed strong technical and communication skills.")
    if "weakness" in q:
        return ("I sometimes take on too much responsibility to ensure quality. "
                "I've learned to delegate effectively and trust my team, which has made me a better collaborator.")
    if any(w in q for w in ["why", "interested", "motivation"]):
        return (f"I'm passionate about {job_title} and {company}'s mission inspires me. "
                f"Your innovative approach aligns with my expertise in {top_skills}, "
                f"and I see tremendous potential to contribute and grow here.")
    if "salary" in q or "compensation" in q:
        return ("Based on my experience and market research, I'm looking for a competitive salary "
                "that reflects my skills and contributions. I'm open to discussing the full compensation package.")
    if "challenge" in q:
        return ("I once led a critical project migration under a tight deadline. "
                "By breaking it into milestones, coordinating cross-functionally, and maintaining clear communication, "
                "we delivered on time with zero downtime — a real team success.")
    if "team" in q or "collaborat" in q:
        return ("I thrive in collaborative environments. I believe in transparent communication, "
                "supporting teammates, and leveraging diverse perspectives to build better solutions.")
    if "remote" in q:
        return ("I'm very comfortable with remote work. I'm self-disciplined, proactive in communication, "
                "and experienced with remote collaboration tools like Slack, Zoom, and GitHub.")
    if "available" in q or "start" in q:
        return "I can start within 2–4 weeks after completing my notice period. I'm flexible and eager to begin."
    return (f"With {experience} in {top_skills}, I'm confident I can bring significant value to {company}. "
            f"I'm detail-oriented, results-driven, and always looking to improve both the product and the team around me.")


def github_api(pat, method, endpoint, data=None):
    headers = {
        "Authorization": f"token {pat}",
        "Accept": "application/vnd.github.v3+json",
        "Content-Type": "application/json"
    }
    url = f"https://api.github.com{endpoint}"
    try:
        r = getattr(requests, method)(url, headers=headers, json=data, timeout=15)
        return r.status_code, r.json() if r.text else {}
    except Exception as e:
        return 0, {"error": str(e)}


def push_file_to_github(pat, owner, repo, path, content, message):
    # Check if file exists (to get sha for update)
    status, existing = github_api(pat, "get", f"/repos/{owner}/{repo}/contents/{path}")
    sha = existing.get("sha") if status == 200 else None
    payload = {
        "message": message,
        "content": base64.b64encode(content.encode()).decode(),
    }
    if sha:
        payload["sha"] = sha
    status, resp = github_api(pat, "put", f"/repos/{owner}/{repo}/contents/{path}", payload)
    return status in (200, 201), resp


# ─── Routes ───────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template_string(HTML_PAGE)


@app.route("/api/jobs/search", methods=["GET"])
def search_jobs():
    title    = request.args.get("title", "Software Engineer")
    location = request.args.get("location", "")
    keywords = request.args.get("keywords", "")
    count    = min(int(request.args.get("count", 20)), 50)

    jobs = []
    # Try real LinkedIn scrape
    try:
        query = f"{title} {keywords}".strip()
        loc_q = location or "worldwide"
        url = (f"https://www.linkedin.com/jobs-guest/jobs/api/seeMoreJobPostings/search"
               f"?keywords={requests.utils.quote(query)}&location={requests.utils.quote(loc_q)}&start=0")
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
            "Accept-Language": "en-US,en;q=0.9"
        }
        r = requests.get(url, headers=headers, timeout=8)
        if r.status_code == 200:
            soup = BeautifulSoup(r.text, "lxml")
            cards = soup.find_all("div", class_=re.compile(r"base-card"))
            for card in cards[:count]:
                try:
                    t = card.find(class_=re.compile(r"base-search-card__title"))
                    c = card.find(class_=re.compile(r"base-search-card__subtitle"))
                    l = card.find(class_=re.compile(r"job-search-card__location"))
                    a = card.find("a", href=True)
                    jid = hashlib.md5((t.text.strip() + c.text.strip()).encode()).hexdigest()[:10]
                    jobs.append({
                        "id": jid,
                        "title": t.text.strip() if t else title,
                        "company": c.text.strip() if c else random.choice(COMPANIES),
                        "location": l.text.strip() if l else location,
                        "url": a["href"] if a else "",
                        "type": random.choice(JOB_TYPES),
                        "salary": random.choice(SALARY_RANGES),
                        "description": random.choice(DESCRIPTIONS),
                        "skills": random.sample(SKILLS_DB, 5),
                        "postedAt": datetime.now().strftime("%Y-%m-%d"),
                        "source": "linkedin",
                        "applied": False,
                        "saved": False,
                        "matchScore": random.randint(60, 99),
                        "easyApply": True
                    })
                except Exception:
                    continue
    except Exception:
        pass

    # Fill remaining with mock data
    remaining = count - len(jobs)
    for i in range(remaining):
        jobs.append(make_job(title, location, i))

    return jsonify({"jobs": jobs, "total": len(jobs), "source": "linkedin+mock"})


@app.route("/api/cv/analyze", methods=["POST"])
def analyze_cv():
    cv_text = ""
    if "file" in request.files:
        f = request.files["file"]
        raw = f.read()
        name = f.filename.lower()
        if name.endswith(".pdf"):
            cv_text = extract_text_from_pdf(raw)
        else:
            cv_text = raw.decode("utf-8", errors="ignore")
    elif request.is_json:
        cv_text = request.json.get("text", "")

    if not cv_text.strip():
        return jsonify({"error": "No CV text found"}), 400

    skills     = extract_skills(cv_text)
    experience = extract_experience(cv_text)
    education  = extract_education(cv_text)
    titles     = suggest_titles(skills, experience)
    summary    = (f"Experienced {experience} professional with expertise in "
                  f"{', '.join(skills[:5])} and a background in {education}.")

    return jsonify({
        "skills": skills,
        "experience": experience,
        "education": education,
        "suggestedTitles": titles,
        "summary": summary,
        "wordCount": len(cv_text.split()),
        "analyzed": True
    })


@app.route("/api/ai/answer", methods=["POST"])
def ai_answer():
    data = request.json or {}
    question  = data.get("question", "Tell me about yourself")
    cv_text   = data.get("cvText", "")
    job_title = data.get("jobTitle", "Software Engineer")
    company   = data.get("company", "the company")
    answer    = smart_answer(question, cv_text, job_title, company)
    return jsonify({"answer": answer, "question": question})


@app.route("/api/jobs/apply", methods=["POST"])
def apply_jobs():
    data       = request.json or {}
    jobs       = data.get("jobs", [])
    cv_text    = data.get("cvText", "")
    min_score  = int(data.get("minScore", 70))
    max_apply  = int(data.get("maxApply", 10))

    results = []
    eligible = [j for j in jobs if j.get("matchScore", 0) >= min_score][:max_apply]
    for job in eligible:
        time.sleep(0.05)  # simulate brief delay
        success = random.random() > 0.15  # 85% success rate
        app_id  = f"APP-{hashlib.md5(job['id'].encode()).hexdigest()[:8].upper()}"
        results.append({
            "jobId": job["id"],
            "jobTitle": job.get("title"),
            "company": job.get("company"),
            "applicationId": app_id,
            "status": "submitted" if success else "failed",
            "appliedAt": datetime.now().isoformat(),
            "coverLetter": smart_answer("Why are you interested?", cv_text,
                                        job.get("title",""), job.get("company",""))
        })

    return jsonify({
        "applied": results,
        "totalApplied": len([r for r in results if r["status"] == "submitted"]),
        "failed": len([r for r in results if r["status"] == "failed"])
    })


@app.route("/api/export/word", methods=["POST"])
def export_word():
    data  = request.json or {}
    jobs  = data.get("jobs", [])
    apps  = data.get("applications", [])
    title = data.get("title", "LinkedIn Job Bot — Export")

    if not DOCX_AVAILABLE:
        # Fallback to RTF
        rtf = "{\\rtf1\\ansi\n{\\b LinkedIn Job Bot Export}\\par\n"
        for j in jobs:
            rtf += f"\\par {j.get('title','')} — {j.get('company','')} ({j.get('location','')})\\par\n"
        rtf += "}"
        buf = io.BytesIO(rtf.encode())
        buf.seek(0)
        return send_file(buf, mimetype="application/rtf",
                         as_attachment=True, download_name="jobs_export.rtf")

    doc = Document()

    # Title
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Jobs: {len(jobs)}  |  Applications: {len(apps)}")
    doc.add_paragraph("─" * 60)

    # Jobs Table
    if jobs:
        doc.add_heading("🔍 Job Listings", 1)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["#","Title","Company","Location","Salary","Match%"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True

        for idx, j in enumerate(jobs, 1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = j.get("title","")
            row[2].text = j.get("company","")
            row[3].text = j.get("location","")
            row[4].text = j.get("salary","")
            row[5].text = f"{j.get('matchScore','')}%"

    # Applications Table
    if apps:
        doc.add_paragraph()
        doc.add_heading("✅ Applications Submitted", 1)
        t2 = doc.add_table(rows=1, cols=4)
        t2.style = "Table Grid"
        hdr2 = t2.rows[0].cells
        for i, h in enumerate(["Job Title","Company","App ID","Status"]):
            hdr2[i].text = h
            hdr2[i].paragraphs[0].runs[0].bold = True
        for a in apps:
            row = t2.add_row().cells
            row[0].text = a.get("jobTitle","")
            row[1].text = a.get("company","")
            row[2].text = a.get("applicationId","")
            row[3].text = a.get("status","")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name="linkedin_jobs_export.docx")


@app.route("/api/export/csv", methods=["POST"])
def export_csv():
    data = request.json or {}
    jobs = data.get("jobs", [])
    buf  = io.StringIO()
    w    = csv.writer(buf)
    w.writerow(["#","Title","Company","Location","Type","Salary","Match%","Posted","URL","Applied"])
    for i, j in enumerate(jobs, 1):
        w.writerow([i, j.get("title"), j.get("company"), j.get("location"),
                    j.get("type"), j.get("salary"), j.get("matchScore"),
                    j.get("postedAt"), j.get("url"), j.get("applied")])
    output = io.BytesIO(buf.getvalue().encode())
    output.seek(0)
    return send_file(output, mimetype="text/csv", as_attachment=True,
                     download_name="jobs_export.csv")


@app.route("/api/export/json", methods=["POST"])
def export_json_file():
    data = request.json or {}
    buf  = io.BytesIO(json.dumps(data, indent=2).encode())
    buf.seek(0)
    return send_file(buf, mimetype="application/json", as_attachment=True,
                     download_name="linkedin_jobs_export.json")


@app.route("/api/github/validate", methods=["POST"])
def github_validate():
    pat = (request.json or {}).get("pat", "")
    if not pat:
        return jsonify({"valid": False, "error": "No PAT provided"}), 400
    status, data = github_api(pat, "get", "/user")
    if status == 200:
        return jsonify({"valid": True, "login": data.get("login"), "name": data.get("name"),
                        "avatar": data.get("avatar_url"), "repos": data.get("public_repos")})
    return jsonify({"valid": False, "error": data.get("message","Invalid token")}), 401


@app.route("/api/github/repos", methods=["POST"])
def github_repos():
    pat = (request.json or {}).get("pat", "")
    status, data = github_api(pat, "get", "/user/repos?per_page=50&sort=updated")
    if status == 200:
        repos = [{"name": r["name"], "full_name": r["full_name"],
                  "private": r["private"], "url": r["html_url"]} for r in data]
        return jsonify({"repos": repos})
    return jsonify({"error": "Failed to fetch repos"}), 400


@app.route("/api/github/create-repo", methods=["POST"])
def github_create_repo():
    data = request.json or {}
    pat  = data.get("pat", "")
    name = data.get("name", "linkedin-job-results")
    desc = data.get("description", "LinkedIn Job Bot results")
    status, resp = github_api(pat, "post", "/user/repos", {
        "name": name, "description": desc,
        "private": data.get("private", False), "auto_init": True
    })
    if status == 201:
        return jsonify({"created": True, "url": resp.get("html_url"),
                        "full_name": resp.get("full_name")})
    return jsonify({"error": resp.get("message","Failed to create repo")}), 400


@app.route("/api/github/push", methods=["POST"])
def github_push():
    data    = request.json or {}
    pat     = data.get("pat", "")
    repo    = data.get("repo", "")        # "owner/repo"
    jobs    = data.get("jobs", [])
    apps    = data.get("applications", [])
    folder  = data.get("folder", "job-bot-results")
    fmt     = data.get("format", "json")  # json | csv | both

    if not pat or not repo:
        return jsonify({"error": "PAT and repo required"}), 400

    owner, repo_name = (repo.split("/") + [""])[:2]
    if not repo_name:
        return jsonify({"error": "Invalid repo format (owner/repo)"}), 400

    ts      = datetime.now().strftime("%Y%m%d_%H%M%S")
    pushed  = []
    errors  = []

    def push(path, content, msg):
        ok, _ = push_file_to_github(pat, owner, repo_name, path, content, msg)
        (pushed if ok else errors).append(path)

    # README
    readme = f"# LinkedIn Job Bot Results\n\nGenerated: {datetime.now().isoformat()}\n\nJobs: {len(jobs)} | Applications: {len(apps)}\n"
    push(f"{folder}/README.md", readme, "🤖 Update README")

    if fmt in ("json","both"):
        push(f"{folder}/jobs_{ts}.json",
             json.dumps(jobs, indent=2),
             f"💼 Add {len(jobs)} jobs ({ts})")
        if apps:
            push(f"{folder}/applications_{ts}.json",
                 json.dumps(apps, indent=2),
                 f"✅ Add {len(apps)} applications ({ts})")

    if fmt in ("csv","both"):
        csv_lines = ["#,Title,Company,Location,Salary,Match%,Posted,URL"]
        for i, j in enumerate(jobs,1):
            csv_lines.append(f"{i},{j.get('title')},{j.get('company')},"
                             f"{j.get('location')},{j.get('salary')},"
                             f"{j.get('matchScore')},{j.get('postedAt')},{j.get('url')}")
        push(f"{folder}/jobs_{ts}.csv", "\n".join(csv_lines),
             f"📊 Add CSV ({ts})")

    # Session summary
    summary = f"""# Session Summary — {datetime.now().strftime('%Y-%m-%d %H:%M')}

| Metric | Value |
|--------|-------|
| Jobs Found | {len(jobs)} |
| Applications | {len(apps)} |
| Files Pushed | {len(pushed)} |
| Timestamp | {ts} |

## Top Jobs
"""
    for j in jobs[:10]:
        summary += f"- **{j.get('title')}** @ {j.get('company')} — {j.get('location')} ({j.get('matchScore')}% match)\n"

    push(f"{folder}/session_{ts}.md", summary, f"📝 Session summary {ts}")

    return jsonify({
        "pushed": pushed,
        "errors": errors,
        "repo": f"https://github.com/{repo}",
        "folder": folder
    })


# ─── Frontend HTML ────────────────────────────────────────────────────────────

HTML_PAGE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1.0"/>
<title>LinkedIn Job Bot 🤖</title>
<script src="https://cdn.tailwindcss.com"></script>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@fortawesome/fontawesome-free@6.4.0/css/all.min.css"/>
<script src="https://cdn.jsdelivr.net/npm/axios@1.6.0/dist/axios.min.js"></script>
<style>
  :root{--primary:#0077b5;--primary-dark:#005f8e;--accent:#00a0dc;--success:#28a745;--danger:#dc3545;--warning:#ffc107;}
  *{box-sizing:border-box;margin:0;padding:0;}
  body{font-family:'Segoe UI',system-ui,sans-serif;background:linear-gradient(135deg,#0077b5 0%,#00a0dc 50%,#004182 100%);min-height:100vh;}

  /* ── Nav ── */
  .navbar{background:rgba(0,0,0,.35);backdrop-filter:blur(10px);padding:.75rem 1rem;position:sticky;top:0;z-index:100;border-bottom:1px solid rgba(255,255,255,.15);}
  .nav-inner{max-width:1200px;margin:0 auto;display:flex;align-items:center;justify-content:space-between;gap:.5rem;flex-wrap:wrap;}
  .brand{color:#fff;font-weight:800;font-size:1.25rem;display:flex;align-items:center;gap:.4rem;white-space:nowrap;}
  .nav-tabs{display:flex;gap:.3rem;flex-wrap:wrap;}
  .nav-tab{background:transparent;border:1px solid rgba(255,255,255,.3);color:rgba(255,255,255,.8);padding:.4rem .8rem;border-radius:6px;cursor:pointer;font-size:.82rem;transition:all .2s;}
  .nav-tab:hover,.nav-tab.active{background:#fff;color:var(--primary);border-color:#fff;font-weight:600;}

  /* ── Layout ── */
  .container{max-width:1200px;margin:0 auto;padding:1rem;}
  .page{display:none;animation:fadeIn .3s ease;}
  .page.active{display:block;}
  @keyframes fadeIn{from{opacity:0;transform:translateY(8px)}to{opacity:1;transform:translateY(0)}}

  /* ── Card ── */
  .card{background:rgba(255,255,255,.97);border-radius:14px;padding:1.25rem;margin-bottom:1rem;box-shadow:0 4px 24px rgba(0,0,0,.12);}
  .card-title{font-size:1.1rem;font-weight:700;color:#1a1a1a;margin-bottom:.75rem;display:flex;align-items:center;gap:.5rem;}

  /* ── Buttons ── */
  .btn{display:inline-flex;align-items:center;gap:.4rem;padding:.55rem 1.1rem;border-radius:8px;border:none;cursor:pointer;font-size:.88rem;font-weight:600;transition:all .2s;white-space:nowrap;}
  .btn-primary{background:var(--primary);color:#fff;} .btn-primary:hover{background:var(--primary-dark);}
  .btn-success{background:var(--success);color:#fff;} .btn-success:hover{background:#218838;}
  .btn-danger{background:var(--danger);color:#fff;}   .btn-danger:hover{background:#c82333;}
  .btn-warning{background:var(--warning);color:#333;} .btn-warning:hover{background:#e0a800;}
  .btn-outline{background:transparent;color:var(--primary);border:2px solid var(--primary);}
  .btn-outline:hover{background:var(--primary);color:#fff;}
  .btn-sm{padding:.35rem .7rem;font-size:.8rem;}
  .btn:disabled{opacity:.5;cursor:not-allowed;}

  /* ── Form ── */
  input,select,textarea{width:100%;padding:.6rem .85rem;border:1.5px solid #d1d5db;border-radius:8px;font-size:.9rem;transition:border-color .2s;background:#fff;}
  input:focus,select:focus,textarea:focus{outline:none;border-color:var(--primary);box-shadow:0 0 0 3px rgba(0,119,181,.12);}
  label{display:block;font-size:.82rem;font-weight:600;color:#374151;margin-bottom:.3rem;}
  .form-group{margin-bottom:.9rem;}
  .grid-2{display:grid;grid-template-columns:1fr 1fr;gap:.75rem;}
  .grid-3{display:grid;grid-template-columns:1fr 1fr 1fr;gap:.75rem;}
  @media(max-width:640px){.grid-2,.grid-3{grid-template-columns:1fr;}}

  /* ── Drop Zone ── */
  .drop-zone{border:2.5px dashed #0077b5;border-radius:12px;padding:2rem 1rem;text-align:center;cursor:pointer;transition:all .25s;background:rgba(0,119,181,.04);}
  .drop-zone:hover,.drop-zone.drag-over{background:rgba(0,119,181,.1);border-color:#00a0dc;}
  .drop-zone i{font-size:2.5rem;color:#0077b5;margin-bottom:.5rem;}

  /* ── Job Card ── */
  .job-card{background:#fff;border-radius:10px;padding:1rem;border:1.5px solid #e5e7eb;transition:all .22s;position:relative;}
  .job-card:hover{border-color:var(--primary);box-shadow:0 4px 16px rgba(0,119,181,.15);}
  .job-card.applied-card{border-color:var(--success);background:rgba(40,167,69,.04);}
  .match-badge{position:absolute;top:.7rem;right:.7rem;padding:.2rem .6rem;border-radius:20px;font-size:.75rem;font-weight:700;}
  .match-high{background:#d4edda;color:#155724;}
  .match-med{background:#fff3cd;color:#856404;}
  .match-low{background:#f8d7da;color:#721c24;}
  .job-title{font-weight:700;font-size:.95rem;color:#1a1a1a;margin-bottom:.15rem;}
  .job-meta{font-size:.8rem;color:#6b7280;display:flex;flex-wrap:wrap;gap:.4rem;margin-bottom:.5rem;}
  .job-meta span{display:flex;align-items:center;gap:.2rem;}
  .skill-tag{display:inline-block;background:#e0f0fa;color:#0077b5;padding:.15rem .55rem;border-radius:20px;font-size:.73rem;font-weight:600;margin:.1rem;}
  .jobs-grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(280px,1fr));gap:.75rem;}
  @media(max-width:480px){.jobs-grid{grid-template-columns:1fr;}}

  /* ── Stats ── */
  .stats-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(120px,1fr));gap:.75rem;}
  .stat-card{background:rgba(255,255,255,.15);border-radius:10px;padding:.9rem;text-align:center;color:#fff;backdrop-filter:blur(5px);}
  .stat-num{font-size:1.8rem;font-weight:800;}
  .stat-label{font-size:.78rem;opacity:.85;}

  /* ── Toast ── */
  #toast-container{position:fixed;top:1rem;right:1rem;z-index:9999;display:flex;flex-direction:column;gap:.5rem;}
  .toast{padding:.75rem 1.1rem;border-radius:10px;color:#fff;font-size:.88rem;font-weight:600;box-shadow:0 4px 20px rgba(0,0,0,.25);display:flex;align-items:center;gap:.5rem;animation:slideIn .3s ease;max-width:320px;}
  .toast-success{background:#28a745;} .toast-error{background:#dc3545;} .toast-info{background:#0077b5;}
  @keyframes slideIn{from{transform:translateX(120%);opacity:0}to{transform:translateX(0);opacity:1}}

  /* ── Progress ── */
  .progress-bar{background:#e9ecef;border-radius:999px;height:10px;overflow:hidden;}
  .progress-fill{height:100%;background:linear-gradient(90deg,#0077b5,#00a0dc);transition:width .6s ease;border-radius:999px;}

  /* ── Log ── */
  .log-box{background:#0f172a;border-radius:10px;padding:1rem;font-family:monospace;font-size:.8rem;max-height:200px;overflow-y:auto;color:#94a3b8;}
  .log-success{color:#4ade80;} .log-error{color:#f87171;} .log-info{color:#60a5fa;}

  /* ── Responsive helpers ── */
  .flex-wrap-gap{display:flex;flex-wrap:wrap;gap:.5rem;align-items:center;}
  .hidden{display:none!important;}
  @media(max-width:480px){.btn{padding:.45rem .75rem;font-size:.82rem;} .card{padding:1rem;}}
</style>
</head>
<body>

<!-- NAVBAR -->
<nav class="navbar">
  <div class="nav-inner">
    <div class="brand"><i class="fab fa-linkedin"></i> LinkedIn Job Bot</div>
    <div class="nav-tabs">
      <button class="nav-tab active" onclick="showPage('search')"><i class="fas fa-search"></i> <span class="hidden sm:inline">Search</span></button>
      <button class="nav-tab" onclick="showPage('cv')"><i class="fas fa-file-alt"></i> <span class="hidden sm:inline">CV</span></button>
      <button class="nav-tab" onclick="showPage('apply')"><i class="fas fa-paper-plane"></i> <span class="hidden sm:inline">Apply</span></button>
      <button class="nav-tab" onclick="showPage('ai')"><i class="fas fa-robot"></i> <span class="hidden sm:inline">AI</span></button>
      <button class="nav-tab" onclick="showPage('github')"><i class="fab fa-github"></i> <span class="hidden sm:inline">GitHub</span></button>
      <button class="nav-tab" onclick="showPage('help')"><i class="fas fa-question-circle"></i> <span class="hidden sm:inline">Help</span></button>
    </div>
  </div>
</nav>

<!-- TOAST -->
<div id="toast-container"></div>

<!-- CONTAINER -->
<div class="container">

  <!-- STATS BAR -->
  <div class="stats-grid" style="margin-bottom:1rem;">
    <div class="stat-card"><div class="stat-num" id="stat-jobs">0</div><div class="stat-label">Jobs Found</div></div>
    <div class="stat-card"><div class="stat-num" id="stat-applied">0</div><div class="stat-label">Applied</div></div>
    <div class="stat-card"><div class="stat-num" id="stat-saved">0</div><div class="stat-label">Saved</div></div>
    <div class="stat-card"><div class="stat-num" id="stat-cv">—</div><div class="stat-label">CV Score</div></div>
  </div>

  <!-- ── PAGE: SEARCH ── -->
  <div id="page-search" class="page active">
    <div class="card">
      <div class="card-title"><i class="fas fa-search" style="color:var(--primary)"></i> Job Search</div>
      <div class="grid-2">
        <div class="form-group"><label>Job Title *</label><input id="s-title" placeholder="e.g. Python Developer" value="Software Engineer"/></div>
        <div class="form-group"><label>Location</label><input id="s-loc" placeholder="e.g. London, Remote"/></div>
      </div>
      <div class="grid-2">
        <div class="form-group"><label>Keywords</label><input id="s-kw" placeholder="e.g. React, Docker, Agile"/></div>
        <div class="form-group"><label>Results (max 50)</label><input id="s-count" type="number" value="20" min="1" max="50"/></div>
      </div>
      <div class="flex-wrap-gap">
        <button class="btn btn-primary" onclick="searchJobs()"><i class="fas fa-search"></i> Search Jobs</button>
        <button class="btn btn-outline" onclick="searchByCv()" id="cv-search-btn" disabled><i class="fas fa-file-alt"></i> Search by CV</button>
        <button class="btn btn-warning btn-sm" onclick="clearJobs()"><i class="fas fa-trash"></i> Clear</button>
      </div>
    </div>

    <!-- Results -->
    <div id="results-area" class="hidden">
      <div class="card" style="padding:.75rem 1.25rem;">
        <div class="flex-wrap-gap" style="justify-content:space-between;">
          <span id="results-count" style="font-weight:700;font-size:.95rem;color:#1a1a1a;"></span>
          <div class="flex-wrap-gap">
            <button class="btn btn-success btn-sm" onclick="exportWord()"><i class="fas fa-file-word"></i> Word</button>
            <button class="btn btn-warning btn-sm" onclick="exportCSV()"><i class="fas fa-file-csv"></i> CSV</button>
            <button class="btn btn-outline btn-sm" onclick="exportJSON()"><i class="fas fa-file-code"></i> JSON</button>
            <button class="btn btn-primary btn-sm" onclick="showPage('apply')"><i class="fas fa-paper-plane"></i> Auto-Apply</button>
          </div>
        </div>
      </div>
      <div class="jobs-grid" id="jobs-grid"></div>
    </div>

    <div id="search-loading" class="hidden card" style="text-align:center;padding:2rem;">
      <i class="fas fa-spinner fa-spin fa-2x" style="color:var(--primary)"></i>
      <p style="margin-top:.75rem;color:#6b7280;font-weight:600;">Fetching jobs from LinkedIn…</p>
    </div>
  </div>

  <!-- ── PAGE: CV ── -->
  <div id="page-cv" class="page">
    <div class="card">
      <div class="card-title"><i class="fas fa-file-alt" style="color:var(--primary)"></i> Upload & Analyze CV</div>
      <div class="drop-zone" id="drop-zone" onclick="document.getElementById('cv-file').click()"
           ondragover="handleDragOver(event)" ondragleave="handleDragLeave(event)" ondrop="handleDrop(event)">
        <i class="fas fa-cloud-upload-alt"></i>
        <p style="font-weight:700;margin-bottom:.25rem;">Drag & Drop your CV here</p>
        <p style="font-size:.82rem;color:#6b7280;">PDF, TXT, or DOC • Click to browse</p>
        <input type="file" id="cv-file" accept=".pdf,.txt,.doc,.docx" style="display:none" onchange="handleFileSelect(event)"/>
      </div>
      <div style="margin-top:1rem;">
        <label>Or paste CV text:</label>
        <textarea id="cv-text" rows="6" placeholder="Paste your CV / resume text here…"></textarea>
      </div>
      <div style="margin-top:.75rem;" class="flex-wrap-gap">
        <button class="btn btn-primary" onclick="analyzeCV()"><i class="fas fa-microscope"></i> Analyze CV</button>
        <button class="btn btn-outline" onclick="clearCV()"><i class="fas fa-times"></i> Clear</button>
      </div>
    </div>

    <div id="cv-results" class="hidden">
      <div class="card">
        <div class="card-title"><i class="fas fa-chart-bar" style="color:var(--success)"></i> CV Analysis Results</div>
        <div class="grid-2" style="margin-bottom:1rem;">
          <div><strong>Experience:</strong> <span id="cv-exp">—</span></div>
          <div><strong>Education:</strong> <span id="cv-edu">—</span></div>
        </div>
        <div style="margin-bottom:.75rem;"><strong>Summary:</strong><p id="cv-summary" style="color:#374151;margin-top:.25rem;"></p></div>
        <div style="margin-bottom:.75rem;"><strong>Skills Detected:</strong><div id="cv-skills" style="margin-top:.4rem;"></div></div>
        <div><strong>Suggested Job Titles:</strong><div id="cv-titles" style="margin-top:.4rem;"></div></div>
      </div>
    </div>
  </div>

  <!-- ── PAGE: APPLY ── -->
  <div id="page-apply" class="page">
    <div class="card">
      <div class="card-title"><i class="fas fa-paper-plane" style="color:var(--primary)"></i> Auto-Apply Settings</div>
      <div class="grid-3">
        <div class="form-group"><label>Min Match Score (%)</label><input id="a-minscore" type="number" value="70" min="0" max="100"/></div>
        <div class="form-group"><label>Max Applications</label><input id="a-maxapply" type="number" value="10" min="1" max="50"/></div>
        <div class="form-group"><label>Delay (ms)</label><input id="a-delay" type="number" value="300" min="0" max="2000"/></div>
      </div>
      <div class="flex-wrap-gap">
        <button class="btn btn-success" onclick="startAutoApply()"><i class="fas fa-rocket"></i> Start Auto-Apply</button>
        <button class="btn btn-danger btn-sm" onclick="stopApply()" id="stop-apply-btn" disabled><i class="fas fa-stop"></i> Stop</button>
      </div>
      <div style="margin-top:1rem;" id="apply-progress-wrap" class="hidden">
        <div style="margin-bottom:.4rem;font-size:.85rem;font-weight:600;" id="apply-prog-label">Applying…</div>
        <div class="progress-bar"><div class="progress-fill" id="apply-prog-bar" style="width:0%"></div></div>
      </div>
    </div>

    <div id="apply-log-card" class="card hidden">
      <div class="card-title"><i class="fas fa-terminal" style="color:var(--primary)"></i> Application Log</div>
      <div class="log-box" id="apply-log"></div>
    </div>

    <div id="apply-results-card" class="card hidden">
      <div class="card-title"><i class="fas fa-check-circle" style="color:var(--success)"></i> Applications Submitted</div>
      <div id="apply-results-list"></div>
    </div>
  </div>

  <!-- ── PAGE: AI ── -->
  <div id="page-ai" class="page">
    <div class="card">
      <div class="card-title"><i class="fas fa-robot" style="color:var(--primary)"></i> AI Interview Answer Generator</div>
      <div class="grid-2">
        <div class="form-group"><label>Job Title</label><input id="ai-title" placeholder="e.g. Senior Python Developer" value="Software Engineer"/></div>
        <div class="form-group"><label>Company</label><input id="ai-company" placeholder="e.g. Google" value="the company"/></div>
      </div>
      <div class="form-group"><label>Your Question</label><textarea id="ai-question" rows="3" placeholder="e.g. Tell me about yourself…"></textarea></div>
      <div class="flex-wrap-gap" style="margin-bottom:.75rem;">
        <button class="btn btn-sm btn-outline" onclick="setQ('Tell me about yourself')">About Me</button>
        <button class="btn btn-sm btn-outline" onclick="setQ('What are your strengths?')">Strengths</button>
        <button class="btn btn-sm btn-outline" onclick="setQ('What is your weakness?')">Weakness</button>
        <button class="btn btn-sm btn-outline" onclick="setQ('Why do you want this job?')">Motivation</button>
        <button class="btn btn-sm btn-outline" onclick="setQ('What are your salary expectations?')">Salary</button>
        <button class="btn btn-sm btn-outline" onclick="setQ('Describe a challenge you faced')">Challenge</button>
      </div>
      <button class="btn btn-primary" onclick="generateAnswer()"><i class="fas fa-magic"></i> Generate Answer</button>
    </div>

    <div id="ai-answer-card" class="card hidden">
      <div class="card-title"><i class="fas fa-comment-dots" style="color:var(--success)"></i> Generated Answer</div>
      <p id="ai-answer-text" style="line-height:1.7;color:#1f2937;"></p>
      <div style="margin-top:.75rem;" class="flex-wrap-gap">
        <button class="btn btn-outline btn-sm" onclick="copyAnswer()"><i class="fas fa-copy"></i> Copy</button>
        <button class="btn btn-success btn-sm" onclick="saveAnswer()"><i class="fas fa-bookmark"></i> Save</button>
      </div>
    </div>

    <div id="answer-history-card" class="card hidden">
      <div class="card-title"><i class="fas fa-history" style="color:var(--primary)"></i> Answer History</div>
      <div id="answer-history-list"></div>
    </div>
  </div>

  <!-- ── PAGE: GITHUB ── -->
  <div id="page-github" class="page">
    <div class="card">
      <div class="card-title"><i class="fab fa-github" style="color:#333"></i> GitHub Integration</div>
      <div class="form-group">
        <label>GitHub Personal Access Token (PAT)</label>
        <div style="display:flex;gap:.5rem;">
          <input id="gh-pat" type="password" placeholder="ghp_xxxxxxxxxxxxxxxxxxxx"/>
          <button class="btn btn-primary btn-sm" onclick="validateGH()" style="white-space:nowrap;"><i class="fas fa-key"></i> Validate</button>
        </div>
      </div>
      <div id="gh-user-card" class="hidden" style="padding:.75rem;background:#f0f9ff;border-radius:8px;margin-bottom:.75rem;display:flex;align-items:center;gap:.75rem;">
        <img id="gh-avatar" src="" style="width:44px;height:44px;border-radius:50%;"/>
        <div><strong id="gh-username"></strong><br/><span id="gh-repocount" style="font-size:.8rem;color:#6b7280;"></span></div>
      </div>

      <div class="grid-2">
        <div class="form-group">
          <label>Target Repository</label>
          <input id="gh-repo" placeholder="owner/repo-name"/>
        </div>
        <div class="form-group">
          <label>Folder in Repo</label>
          <input id="gh-folder" value="job-bot-results"/>
        </div>
      </div>
      <div class="form-group">
        <label>Save Format</label>
        <select id="gh-format">
          <option value="json">JSON only</option>
          <option value="csv">CSV only</option>
          <option value="both" selected>Both JSON + CSV</option>
        </select>
      </div>
      <div class="flex-wrap-gap">
        <button class="btn btn-primary" onclick="pushToGH()"><i class="fas fa-upload"></i> Push Results to GitHub</button>
        <button class="btn btn-outline btn-sm" onclick="loadRepos()"><i class="fas fa-list"></i> Load My Repos</button>
        <button class="btn btn-success btn-sm" onclick="showCreateRepo()"><i class="fas fa-plus"></i> Create Repo</button>
      </div>
    </div>

    <!-- Create repo -->
    <div id="create-repo-card" class="card hidden">
      <div class="card-title"><i class="fas fa-plus-circle" style="color:var(--success)"></i> Create New Repository</div>
      <div class="grid-2">
        <div class="form-group"><label>Repo Name</label><input id="new-repo-name" value="linkedin-job-results"/></div>
        <div class="form-group"><label>Description</label><input id="new-repo-desc" value="LinkedIn Job Bot results"/></div>
      </div>
      <button class="btn btn-success" onclick="createRepo()"><i class="fas fa-check"></i> Create Repository</button>
    </div>

    <!-- Repos list -->
    <div id="repos-card" class="card hidden">
      <div class="card-title"><i class="fas fa-code-branch" style="color:var(--primary)"></i> Your Repositories</div>
      <div id="repos-list"></div>
    </div>

    <!-- Sync log -->
    <div id="gh-log-card" class="card hidden">
      <div class="card-title"><i class="fas fa-terminal"></i> Sync Log</div>
      <div class="log-box" id="gh-log"></div>
    </div>
  </div>

  <!-- ── PAGE: HELP ── -->
  <div id="page-help" class="page">
    <div class="card">
      <div class="card-title"><i class="fas fa-question-circle" style="color:var(--primary)"></i> How to Use LinkedIn Job Bot</div>
      <div style="line-height:1.8;color:#374151;">
        <h3 style="color:var(--primary);margin-bottom:.5rem;">🚀 Quick Start</h3>
        <ol style="padding-left:1.2rem;margin-bottom:1rem;">
          <li><strong>Search</strong> — Enter a job title &amp; location, click <em>Search Jobs</em></li>
          <li><strong>Upload CV</strong> — Drag &amp; drop or paste your CV, click <em>Analyze CV</em></li>
          <li><strong>Auto-Apply</strong> — Set min score &amp; max applications, click <em>Start Auto-Apply</em></li>
          <li><strong>AI Answers</strong> — Type any interview question, get a smart answer instantly</li>
          <li><strong>Export</strong> — Download results as Word, CSV, or JSON</li>
          <li><strong>GitHub</strong> — Add your PAT, select a repo, and push all results</li>
        </ol>

        <h3 style="color:var(--primary);margin-bottom:.5rem;">💻 Run Locally (Step by Step)</h3>
        <div style="background:#0f172a;color:#94a3b8;padding:1rem;border-radius:8px;font-family:monospace;font-size:.82rem;margin-bottom:1rem;overflow-x:auto;">
          <span style="color:#4ade80;"># 1. Install Python 3.10+ from https://python.org</span><br/>
          <span style="color:#4ade80;"># 2. Open Terminal / Command Prompt</span><br/>
          git clone https://github.com/abdullah-v-cmd/linkedin-job-bot.git<br/>
          cd linkedin-job-bot<br/>
          pip install flask flask-cors requests beautifulsoup4 python-docx PyPDF2 lxml<br/>
          python app.py<br/>
          <span style="color:#4ade80;"># 3. Open http://localhost:5000 in your browser</span>
        </div>

        <h3 style="color:var(--primary);margin-bottom:.5rem;">📁 Export Formats</h3>
        <ul style="padding-left:1.2rem;margin-bottom:1rem;">
          <li><strong>Word (.docx)</strong> — Professional table with all job details + applications</li>
          <li><strong>CSV</strong> — Spreadsheet-ready, opens in Excel / Google Sheets</li>
          <li><strong>JSON</strong> — Raw data for developers or further processing</li>
        </ul>

        <h3 style="color:var(--primary);margin-bottom:.5rem;">🔑 GitHub PAT Setup</h3>
        <ol style="padding-left:1.2rem;margin-bottom:1rem;">
          <li>Go to <a href="https://github.com/settings/tokens" target="_blank" style="color:var(--primary)">github.com/settings/tokens</a></li>
          <li>Click <em>Generate new token (classic)</em></li>
          <li>Select scopes: <code>repo</code>, <code>read:user</code></li>
          <li>Copy token → paste in the GitHub tab here</li>
        </ol>

        <h3 style="color:var(--primary);margin-bottom:.5rem;">⚠️ Notes</h3>
        <ul style="padding-left:1.2rem;">
          <li>LinkedIn may block scraping — the bot falls back to realistic mock data</li>
          <li>Auto-apply simulates the process (real LinkedIn apply requires browser automation)</li>
          <li>Keep your PAT secret — never share it publicly</li>
        </ul>
      </div>
    </div>
  </div>

</div><!-- /container -->

<script>
// ── State ──────────────────────────────────────────────────────────────────
const S = {
  jobs: [], applications: [], cvText: '', cvData: null,
  ghPat: '', ghUser: null, answerHistory: [], stopApplyFlag: false
};

// ── Helpers ────────────────────────────────────────────────────────────────
function showPage(id) {
  document.querySelectorAll('.page').forEach(p => p.classList.remove('active'));
  document.querySelectorAll('.nav-tab').forEach(t => t.classList.remove('active'));
  document.getElementById('page-' + id).classList.add('active');
  const tabs = {search:0,cv:1,apply:2,ai:3,github:4,help:5};
  document.querySelectorAll('.nav-tab')[tabs[id]]?.classList.add('active');
}
function toast(msg, type='info') {
  const c = document.getElementById('toast-container');
  const el = document.createElement('div');
  el.className = `toast toast-${type}`;
  el.innerHTML = `<i class="fas fa-${type==='success'?'check-circle':type==='error'?'exclamation-circle':'info-circle'}"></i> ${msg}`;
  c.appendChild(el);
  setTimeout(() => el.remove(), 4000);
}
function updateStats() {
  document.getElementById('stat-jobs').textContent = S.jobs.length;
  document.getElementById('stat-applied').textContent = S.applications.length;
  document.getElementById('stat-saved').textContent = S.jobs.filter(j=>j.saved).length;
  document.getElementById('stat-cv').textContent = S.cvData ? S.cvData.skills.length + ' skills' : '—';
}
function matchClass(s) { return s>=80?'match-high':s>=60?'match-med':'match-low'; }
function setQ(q) { document.getElementById('ai-question').value = q; }

// ── Search ─────────────────────────────────────────────────────────────────
async function searchJobs() {
  const title = document.getElementById('s-title').value.trim();
  if (!title) { toast('Enter a job title first','error'); return; }
  document.getElementById('search-loading').classList.remove('hidden');
  document.getElementById('results-area').classList.add('hidden');
  try {
    const r = await axios.get('/api/jobs/search', {params:{
      title, location: document.getElementById('s-loc').value,
      keywords: document.getElementById('s-kw').value,
      count: document.getElementById('s-count').value
    }});
    S.jobs = r.data.jobs;
    renderJobs();
    toast(`Found ${S.jobs.length} jobs!`, 'success');
  } catch(e) { toast('Search failed: ' + (e.response?.data?.error || e.message), 'error'); }
  document.getElementById('search-loading').classList.add('hidden');
}

function searchByCv() {
  if (!S.cvData) { toast('Analyze your CV first','error'); return; }
  document.getElementById('s-title').value = S.cvData.suggestedTitles[0] || 'Software Engineer';
  searchJobs();
}

function renderJobs() {
  const grid = document.getElementById('jobs-grid');
  grid.innerHTML = '';
  S.jobs.forEach((j, i) => {
    const applied = j.applied ? 'applied-card' : '';
    const mc = matchClass(j.matchScore);
    grid.innerHTML += `
    <div class="job-card ${applied}" id="jcard-${j.id}">
      <span class="match-badge ${mc}">${j.matchScore}%</span>
      <div class="job-title">${j.title}</div>
      <div class="job-meta">
        <span><i class="fas fa-building"></i>${j.company}</span>
        <span><i class="fas fa-map-marker-alt"></i>${j.location}</span>
        <span><i class="fas fa-briefcase"></i>${j.type}</span>
        <span><i class="fas fa-dollar-sign"></i>${j.salary}</span>
        <span><i class="fas fa-calendar"></i>${j.postedAt}</span>
      </div>
      <div style="margin-bottom:.6rem;">${(j.skills||[]).slice(0,4).map(s=>`<span class="skill-tag">${s}</span>`).join('')}</div>
      <div class="flex-wrap-gap">
        <a href="${j.url}" target="_blank" class="btn btn-outline btn-sm"><i class="fas fa-external-link-alt"></i> View</a>
        <button class="btn btn-primary btn-sm" onclick="quickApply('${j.id}')"><i class="fas fa-paper-plane"></i> Apply</button>
        <button class="btn btn-sm ${j.saved?'btn-warning':'btn-outline'}" onclick="toggleSave('${j.id}')"><i class="fas fa-bookmark"></i></button>
      </div>
    </div>`;
  });
  document.getElementById('results-count').textContent = `${S.jobs.length} jobs found`;
  document.getElementById('results-area').classList.remove('hidden');
  updateStats();
}

async function quickApply(id) {
  const job = S.jobs.find(j=>j.id===id); if (!job||job.applied) return;
  try {
    const r = await axios.post('/api/jobs/apply', {jobs:[job], cvText:S.cvText, minScore:0, maxApply:1});
    if (r.data.applied?.length) {
      job.applied = true;
      S.applications.push(...r.data.applied);
      document.getElementById('jcard-'+id)?.classList.add('applied-card');
      toast(`Applied to ${job.title} @ ${job.company}!`, 'success');
      updateStats();
    }
  } catch(e) { toast('Apply failed','error'); }
}

function toggleSave(id) {
  const j = S.jobs.find(j=>j.id===id); if (!j) return;
  j.saved = !j.saved;
  renderJobs();
  toast(j.saved ? 'Job saved!' : 'Job unsaved', 'info');
}

function clearJobs() { if(confirm('Clear all job results?')){S.jobs=[];document.getElementById('results-area').classList.add('hidden');updateStats();} }

// ── CV ─────────────────────────────────────────────────────────────────────
function handleDragOver(e) { e.preventDefault(); document.getElementById('drop-zone').classList.add('drag-over'); }
function handleDragLeave(e) { document.getElementById('drop-zone').classList.remove('drag-over'); }
function handleDrop(e) {
  e.preventDefault(); handleDragLeave(e);
  const f = e.dataTransfer.files[0]; if(f) readFile(f);
}
function handleFileSelect(e) { const f = e.target.files[0]; if(f) readFile(f); }
function readFile(f) {
  if (f.name.endsWith('.pdf')) {
    const fd = new FormData(); fd.append('file', f);
    axios.post('/api/cv/analyze', fd, {headers:{'Content-Type':'multipart/form-data'}})
      .then(r => { showCvResults(r.data); toast('CV analyzed!','success'); })
      .catch(() => toast('PDF analysis failed','error'));
    return;
  }
  const reader = new FileReader();
  reader.onload = e => {
    document.getElementById('cv-text').value = e.target.result;
    toast(`Loaded: ${f.name}`, 'info');
  };
  reader.readAsText(f);
}

async function analyzeCV() {
  const text = document.getElementById('cv-text').value.trim();
  if (!text) { toast('Please upload or paste your CV first','error'); return; }
  S.cvText = text;
  try {
    const r = await axios.post('/api/cv/analyze', {text});
    showCvResults(r.data);
    toast('CV analyzed!','success');
  } catch(e) { toast('Analysis failed','error'); }
}

function showCvResults(data) {
  S.cvData = data; S.cvText = S.cvText || document.getElementById('cv-text').value;
  document.getElementById('cv-exp').textContent = data.experience;
  document.getElementById('cv-edu').textContent = data.education;
  document.getElementById('cv-summary').textContent = data.summary;
  document.getElementById('cv-skills').innerHTML = data.skills.map(s=>`<span class="skill-tag">${s}</span>`).join('');
  document.getElementById('cv-titles').innerHTML = data.suggestedTitles.map(t=>`<span class="skill-tag" style="background:#d4edda;color:#155724;">${t}</span>`).join('');
  document.getElementById('cv-results').classList.remove('hidden');
  document.getElementById('cv-search-btn').disabled = false;
  updateStats();
}

function clearCV() {
  document.getElementById('cv-text').value='';
  document.getElementById('cv-results').classList.add('hidden');
  document.getElementById('cv-file').value='';
  S.cvText=''; S.cvData=null;
  document.getElementById('cv-search-btn').disabled=true;
  updateStats();
}

// ── Apply ──────────────────────────────────────────────────────────────────
async function startAutoApply() {
  if (!S.jobs.length) { toast('Search for jobs first!','error'); return; }
  const minScore = parseInt(document.getElementById('a-minscore').value)||70;
  const maxApply = parseInt(document.getElementById('a-maxapply').value)||10;
  const eligible = S.jobs.filter(j=>!j.applied && j.matchScore>=minScore);
  if (!eligible.length) { toast('No eligible jobs. Lower the min score?','error'); return; }

  S.stopApplyFlag = false;
  document.getElementById('stop-apply-btn').disabled = false;
  document.getElementById('apply-progress-wrap').classList.remove('hidden');
  document.getElementById('apply-log-card').classList.remove('hidden');
  const logEl = document.getElementById('apply-log');
  logEl.innerHTML = '';

  const logLine = (msg, cls='') => {
    const span = document.createElement('span');
    span.className = cls; span.textContent = `[${new Date().toLocaleTimeString()}] ${msg}\n`;
    logEl.appendChild(span); logEl.scrollTop = logEl.scrollHeight;
  };

  logLine(`🚀 Starting auto-apply: ${Math.min(eligible.length,maxApply)} jobs (min ${minScore}% match)`, 'log-info');
  const toApply = eligible.slice(0, maxApply);
  let done = 0;

  try {
    const r = await axios.post('/api/jobs/apply', {jobs: toApply, cvText: S.cvText, minScore, maxApply});
    for (const res of (r.data.applied || [])) {
      if (S.stopApplyFlag) { logLine('🛑 Stopped by user','log-error'); break; }
      const job = S.jobs.find(j=>j.id===res.jobId);
      if (job) job.applied = true;
      S.applications.push(res);
      done++;
      const pct = Math.round(done/toApply.length*100);
      document.getElementById('apply-prog-bar').style.width = pct+'%';
      document.getElementById('apply-prog-label').textContent = `Applying… ${done}/${toApply.length}`;
      const ok = res.status==='submitted';
      logLine(`${ok?'✅':'❌'} ${res.jobTitle} @ ${res.company} — ${res.status}`, ok?'log-success':'log-error');
      await new Promise(res=>setTimeout(res, parseInt(document.getElementById('a-delay').value)||300));
    }
    renderApplyResults();
    toast(`Auto-apply done! ${r.data.totalApplied} submitted.`, 'success');
  } catch(e) { logLine('❌ Error: '+e.message,'log-error'); toast('Apply error','error'); }

  document.getElementById('stop-apply-btn').disabled = true;
  updateStats();
  if (S.jobs.length) renderJobs();
}

function stopApply() { S.stopApplyFlag=true; document.getElementById('stop-apply-btn').disabled=true; }

function renderApplyResults() {
  const card = document.getElementById('apply-results-card');
  const list = document.getElementById('apply-results-list');
  list.innerHTML = S.applications.slice(-20).map(a=>
    `<div style="padding:.6rem;border-radius:8px;background:${a.status==='submitted'?'#d4edda':'#f8d7da'};margin-bottom:.4rem;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:.3rem;">
      <div><strong>${a.jobTitle}</strong> @ ${a.company}</div>
      <div style="font-size:.8rem;color:#6b7280;">${a.applicationId} • ${a.status}</div>
    </div>`
  ).join('');
  card.classList.remove('hidden');
}

// ── AI ─────────────────────────────────────────────────────────────────────
async function generateAnswer() {
  const q = document.getElementById('ai-question').value.trim();
  if (!q) { toast('Enter a question first','error'); return; }
  try {
    const r = await axios.post('/api/ai/answer', {
      question: q,
      cvText: S.cvText,
      jobTitle: document.getElementById('ai-title').value,
      company: document.getElementById('ai-company').value
    });
    document.getElementById('ai-answer-text').textContent = r.data.answer;
    document.getElementById('ai-answer-card').classList.remove('hidden');
    toast('Answer generated!','success');
  } catch(e) { toast('Error generating answer','error'); }
}

function copyAnswer() {
  navigator.clipboard.writeText(document.getElementById('ai-answer-text').textContent);
  toast('Copied to clipboard!','success');
}

function saveAnswer() {
  S.answerHistory.push({
    question: document.getElementById('ai-question').value,
    answer: document.getElementById('ai-answer-text').textContent,
    savedAt: new Date().toLocaleString()
  });
  renderAnswerHistory();
  toast('Answer saved!','success');
}

function renderAnswerHistory() {
  if (!S.answerHistory.length) return;
  const card = document.getElementById('answer-history-card');
  const list = document.getElementById('answer-history-list');
  list.innerHTML = S.answerHistory.slice(-5).reverse().map(a=>
    `<div style="padding:.75rem;background:#f8fafc;border-radius:8px;margin-bottom:.5rem;border-left:3px solid var(--primary);">
      <strong style="font-size:.85rem;color:var(--primary);">${a.question}</strong>
      <p style="font-size:.85rem;color:#374151;margin-top:.3rem;">${a.answer}</p>
      <small style="color:#9ca3af;">${a.savedAt}</small>
    </div>`
  ).join('');
  card.classList.remove('hidden');
}

// ── GitHub ──────────────────────────────────────────────────────────────────
async function validateGH() {
  const pat = document.getElementById('gh-pat').value.trim();
  if (!pat) { toast('Enter a PAT first','error'); return; }
  try {
    const r = await axios.post('/api/github/validate', {pat});
    S.ghPat = pat; S.ghUser = r.data;
    document.getElementById('gh-avatar').src = r.data.avatar;
    document.getElementById('gh-username').textContent = r.data.name || r.data.login;
    document.getElementById('gh-repocount').textContent = `@${r.data.login} • ${r.data.repos} repos`;
    document.getElementById('gh-repo').value = r.data.login + '/linkedin-job-results';
    document.getElementById('gh-user-card').style.display = 'flex';
    document.getElementById('gh-user-card').classList.remove('hidden');
    toast('GitHub connected!','success');
  } catch(e) { toast('Invalid token: '+( e.response?.data?.error||e.message),'error'); }
}

async function loadRepos() {
  if (!S.ghPat) { toast('Validate your PAT first','error'); return; }
  try {
    const r = await axios.post('/api/github/repos', {pat:S.ghPat});
    const card = document.getElementById('repos-card');
    document.getElementById('repos-list').innerHTML = r.data.repos.map(repo=>
      `<div style="padding:.5rem .75rem;border-radius:8px;background:#f8fafc;margin-bottom:.3rem;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:.3rem;">
        <div><i class="fas fa-code-branch" style="color:var(--primary);margin-right:.3rem;"></i><strong>${repo.name}</strong>${repo.private?' 🔒':''}</div>
        <button class="btn btn-outline btn-sm" onclick="selectRepo('${repo.full_name}')">Select</button>
      </div>`
    ).join('');
    card.classList.remove('hidden');
  } catch(e) { toast('Failed to load repos','error'); }
}

function selectRepo(fullName) {
  document.getElementById('gh-repo').value = fullName;
  toast(`Selected: ${fullName}`,'success');
}

function showCreateRepo() { document.getElementById('create-repo-card').classList.toggle('hidden'); }

async function createRepo() {
  if (!S.ghPat) { toast('Validate your PAT first','error'); return; }
  const name = document.getElementById('new-repo-name').value.trim();
  const desc = document.getElementById('new-repo-desc').value.trim();
  try {
    const r = await axios.post('/api/github/create-repo', {pat:S.ghPat, name, description:desc});
    document.getElementById('gh-repo').value = r.data.full_name;
    toast(`Repo created: ${r.data.full_name}`,'success');
    document.getElementById('create-repo-card').classList.add('hidden');
  } catch(e) { toast('Create failed: '+(e.response?.data?.error||e.message),'error'); }
}

function ghLog(msg) {
  const el = document.getElementById('gh-log');
  document.getElementById('gh-log-card').classList.remove('hidden');
  const ok = msg.includes('✅')||msg.includes('success');
  const cls = ok?'log-success':msg.includes('❌')?'log-error':'log-info';
  el.innerHTML += `<span class="${cls}">[${new Date().toLocaleTimeString()}] ${msg}\n</span>`;
  el.scrollTop = el.scrollHeight;
}

async function pushToGH() {
  const pat = document.getElementById('gh-pat').value.trim() || S.ghPat;
  const repo = document.getElementById('gh-repo').value.trim();
  if (!pat||!repo) { toast('PAT and repo required','error'); return; }
  if (!S.jobs.length && !S.applications.length) { toast('Nothing to push yet — search for jobs first!','error'); return; }
  ghLog('🔄 Pushing to GitHub…');
  try {
    const r = await axios.post('/api/github/push', {
      pat, repo,
      jobs: S.jobs,
      applications: S.applications,
      folder: document.getElementById('gh-folder').value || 'job-bot-results',
      format: document.getElementById('gh-format').value
    });
    r.data.pushed.forEach(p => ghLog(`✅ Pushed: ${p}`));
    r.data.errors.forEach(p => ghLog(`❌ Failed: ${p}`));
    toast(`Pushed ${r.data.pushed.length} files to GitHub!`,'success');
  } catch(e) { ghLog('❌ Error: '+( e.response?.data?.error||e.message)); toast('Push failed','error'); }
}

// ── Export ─────────────────────────────────────────────────────────────────
function downloadBlob(blob, name) {
  const url = URL.createObjectURL(blob);
  const a = document.createElement('a'); a.href=url; a.download=name; a.click();
  URL.revokeObjectURL(url);
}

async function exportWord() {
  if (!S.jobs.length) { toast('No jobs to export','error'); return; }
  try {
    const r = await axios.post('/api/export/word', {jobs:S.jobs, applications:S.applications, title:'LinkedIn Job Bot Export'}, {responseType:'blob'});
    downloadBlob(r.data, 'jobs_export.docx');
    toast('Word document downloaded!','success');
  } catch(e) { toast('Export failed','error'); }
}

async function exportCSV() {
  if (!S.jobs.length) { toast('No jobs to export','error'); return; }
  try {
    const r = await axios.post('/api/export/csv', {jobs:S.jobs}, {responseType:'blob'});
    downloadBlob(r.data, 'jobs_export.csv');
    toast('CSV downloaded!','success');
  } catch(e) { toast('Export failed','error'); }
}

async function exportJSON() {
  if (!S.jobs.length) { toast('No jobs to export','error'); return; }
  try {
    const r = await axios.post('/api/export/json', {jobs:S.jobs, applications:S.applications}, {responseType:'blob'});
    downloadBlob(r.data, 'jobs_export.json');
    toast('JSON downloaded!','success');
  } catch(e) { toast('Export failed','error'); }
}
</script>
</body>
</html>"""

if __name__ == "__main__":
    print("""
╔══════════════════════════════════════════════╗
║      🤖  LinkedIn Job Bot — Python           ║
║  Mobile-Responsive | Word Export | GitHub    ║
╚══════════════════════════════════════════════╝
  → Running at  http://localhost:5000
  → Ctrl+C to stop
""")
    app.run(host="0.0.0.0", port=5000, debug=False)
